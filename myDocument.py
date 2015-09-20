# coding=utf-8
__author__ = 'yxie'

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches
import os
from sqlalchemy import create_engine
import sqlsoup

basedir = os.path.abspath(os.path.dirname(__file__))
engine = create_engine('sqlite:///' + os.path.join(basedir, 'bible.db'))
db = sqlsoup.SQLSoup(engine)

character = [u'一', u'二', u'三', u'四', u'五', u'六', u'七', u'八', u'九', u'十', u'十一', u'十二', u'十三', u'十四',
             u'十五', u'十六', u'十七', u'十八', u'十九', u'二十', u'二十一', u'二十二', u'二十三', u'二十四', u'二十五',
             u'二十六', u'二十七', u'二十八', u'二十九', u'三十', u'三十一', u'三十二', u'三十三', u'三十四', u'三十五',
             u'三十六', u'三十七', u'三十八', u'三十九', u'四十', u'四十一', u'四十二', u'四十三', u'四十四', u'四十五',
             u'四十六', u'四十七', u'四十八', u'四十九', u'五十', u'五十一', u'五十二', u'五十三', u'五十四', u'五十五',
             u'五十六', u'五十七', u'五十八', u'五十九', u'六十', u'六十一', u'六十二', u'六十三', u'六十四', u'六十五',
             u'六十六']

for book_num in range(1, 40):
    # document = Document()
    book_id = db.BibleID.filter(db.BibleID.SN == book_num).first()
    book = db.Bible.filter(db.Bible.VolumeSN == book_num).all()
    chapter_num = book_id.ChapterNumber
    print chapter_num

    document = Document()
    document.add_heading(book_id.FullName, 0)

    # current_chapter = 1
    # if chapter_num > 1:
    #     document.add_heading(u'第' + character[current_chapter-1] + u'章', 2)
    # p = document.add_paragraph()
    # for verse in book:
    #     if verse.ChapterSN != current_chapter:
    #         current_chapter += 1
    #         document.add_heading(u'第' + character[current_chapter-1] + u'章', 2)
    #         p = document.add_paragraph()
    #     p.add_run(str(verse.VerseSN)).font.superscript = True
    #     run = p.add_run(verse.Lection)
    #     run.font.name = u"Adobe 楷体 Std R"
    #     r = run._element
    #     r.rPr.rFonts.set(qn('w:eastAsia'), u'Adobe 楷体 Std R')

    current_chapter = 1
    p = document.add_paragraph()
    if chapter_num > 1:
        p.add_run(str(current_chapter)+' ').bold = True
    for verse in book:
        if verse.ChapterSN != current_chapter:
            current_chapter += 1
            # document.add_heading(u'第' + character[current_chapter-1] + u'章', 2)
            p = document.add_paragraph()
            p.add_run(str(current_chapter)+' ').bold = True
        p.add_run(str(verse.VerseSN)).font.superscript = True
        run = p.add_run(verse.Lection)
        run.font.name = u"Adobe 楷体 Std R"
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'Adobe 楷体 Std R')

    document.save(book_id.FullName+'.docx')
